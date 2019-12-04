# -*- coding: utf-8 -*-
"""Работа с документами"""

from docx import Document

# Создание нового документа
DOCUMENT = Document()

# Добавление заголовка документа наивысшего уровня
DOCUMENT.add_heading('Document Title', 0)

# Добавляем параграф и заполняем его в переменную p
PARAGRAPH = DOCUMENT.add_paragraph('A plan paragraph having some')

# Добавление параграфа текстом с разными стилями оформления
PARAGRAPH.add_run(' bold').bold = True
PARAGRAPH.add_run(' and some ')
PARAGRAPH.add_run('italic.').italic = True

# Добавление заголовка уровнем ниже
DOCUMENT.add_heading('Heading, level 1', level=1)
DOCUMENT.add_paragraph('Intense quote', style='IntenseQuote')

# Добавление парагрофов со стилем/списком
# неупорядоченный
DOCUMENT.add_paragraph('first item in unordered list', style='ListBullet')
# упорядоченный
DOCUMENT.add_paragraph('first item in order list', style='ListNumber')

# Добавление таблицы (1 строка и 3 столбца)
TABLE = DOCUMENT.add_table(rows=1, cols=3)
HDR_CELLS = TABLE.rows[0].cells
HDR_CELLS[0].text = 'Qty'
HDR_CELLS[1].text = 'Id'
HDR_CELLS[2].text = 'Desc'

# Сохранение документа в файл
DOCUMENT.save('demo.docx')
