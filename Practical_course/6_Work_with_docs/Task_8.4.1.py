# -*- coding: utf-8 -*-
"""Работа с документами doc"""

import os.path
import random
from docx import Document
from openpyxl import Workbook, load_workbook, utils
import PyPDF2


def string_replacement(file_name, str1, str2):
    """
    Функция получает на вход имя файла(документ Word) и две строки.
    Выполняет замену первой найденной строки на другую.
    При успешной замене возвращает True, иначе False
    """
    # Проверка наличия заданного файла
    if os.path.isfile(file_name):
        try:
            document = Document(file_name)
            # Поиск строки в параграфах
            for per in document.paragraphs:
                if per.text == str1:
                    per.text = str2
                    document.save(file_name)
                    return True
            # Поиск строки в таблицах
            for selected_table in document.tables:
                for selected_row in selected_table.rows:
                    for selected_cell in selected_row.cells:
                        if selected_cell.text == str1:
                            selected_cell.text = str2
                            document.save(file_name)
                            return True
        except ValueError:
            print('Неверный фsdfормат файла!')
    return False


def file_square(file_name):
    """
    Функция открывает excel-файл по заданному имени.
    На случайном листе заполняет квадрат 12*12 ячеек случайными значениями
    При успешном выполнении возвращает True иначе False
    """
    # Проверка наличия заданного файла
    if os.path.isfile(file_name):
        try:
            excel_file = load_workbook(file_name)
            # Выбор случайного листа
            selected_sheet = excel_file[random.choice(excel_file.sheetnames)]
            # заполнение ячеек случайными  числами
            for x in range(1, 13):
                for y in range(1, 13):
                    selected_sheet.cell(
                        row=x, column=y, value=random.randint(
                            0, 101))
            excel_file.save(filename=file_name)
            return True
        # except ValueError:
        except utils.exceptions.InvalidFileException:
            print('Не поддерживается формат файла!')
    return False


def read_PDF(file_name, page):
    """
    Функция принимает имя файла и номер страницы.
    Если файл найден и страница указана корректно в консоль выводится считанная информация
    """
    if os.path.isfile(file_name):
        with open(file_name, 'rb') as file:
            read_pdf = PyPDF2.PdfFileReader(file)
            number_pages = read_pdf.getNumPages()
            if page <= number_pages:
                page_content = read_pdf.getPage(page - 1).extractText()
                return page_content.encode('utf-8')
